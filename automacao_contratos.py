import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from datetime import datetime


inf1 = pd.read_excel('arquivo1.xlsx') # Tabela exemplo caso queira incluir

df = pd.read_excel('planilha_que_quero_ler.xlsx', sheet_name='aba do excel desejada') # Planilha excel

doc = Document("Contrato_Modelo.docx")#Modelo base

name_var1 = 'NOME COMPLETO DO FULANO'
name_var2 = 'NOME COMPLETO DO CICLANO'

if name_var1 in df.any():
    texto = 'aqui se coloca o texto que você quer que seja dinamico e que mude caso for a var1'
    assinatura = 'NOME COMPLETO DO FULANO \n Caso queira incluir outra classificação aqui como c.p.f, profissão, etc...'
elif name_var2 in df.values:
    texto = 'aqui se coloca o texto que você quer que seja dinamico e que mude caso for a var2'
    assinatura = 'NOME COMPLETO DO CICLANO \n Caso queira incluir outra classificação aqui como c.p.f, profissão, etc...'
else:
    texto = 'Nenhum nome foi encontrado no Excel'
    assinatura = 'Nenhum nome foi encontrado no Excel'
    
    
for p in doc.paragraphs:
    
    # verificar se o parágrafo contém o texto que deve ser alterado
    if '(texto para alterar dinamicamente)' in p.text:
        # substituir o texto com o nome correspondente ao nome encontrado
        p.text = p.text.replace('(texto para alterar dinamicamente)', texto)
        
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

texto_antigo = 'Fernandópolis-SP, 21 de dezembro de 2022.' # assim estava no modelo que utilizei, mas você pode fazer do jeito que preferir
paragrafo = None
for p in doc.paragraphs: # Aqui busca a var texto antigo no modelo...
    if texto_antigo in p.text:
        paragrafo = p
        break
    
hoje = datetime.today() # add a data para colocar dinamicamente
data = hoje.strftime('%d de %B de %Y') # Como é modelo de contrato, o mês geralmente é nominal, caso queira de outra forma, busque sobre a lib datetime (https://docs.python.org/3/library/datetime.html)

# Dicionário para tradução dos meses, pois a lib está apenas em inglês :\
meses = {'January': 'Janeiro', 'February': 'Fevereiro', 'March': 'Março', 'April': 'Abril', 'May': 'Maio', 'June': 'Junho',
         'July': 'Julho', 'August': 'Agosto', 'September': 'Setembro', 'October': 'Outubro', 'November': 'Novembro', 'December': 'Dezembro'}

if paragrafo is not None: # Colocando a data dinamicamente
    novo_texto = paragrafo.text.replace(texto_antigo, f'Fernandópolis-SP, {data}')
    paragrafo.text = novo_texto
    
for p in doc.paragraphs: #Alterando o nome em inglês do mês para Português
    for mes in meses:
        if mes in p.text:
            p.text = p.text.replace(mes, meses[mes])
            
for p in doc.paragraphs: # Coloquei essa var dentro do meu arquivo word para mudar a assinatura dinamicamente
    if 'VAR_ASS' in p.text:
        p.text = p.text.replace('VAR_ASS', assinatura)

# Colocando as colunas do excel em variaveis de listas
columns1 = list(inf1.columns)

table1 = doc.add_table(rows = 1, cols = len(columns1), style = "Table Grid") # Adicionando modelo de Tabela no Word
table1.autofit = True

for col in range(len(columns1)): # Preenchendo colunas que coloquei no word com a tabela que peguei do excel
    table1.cell(0, col).text = columns1[col]
    
for i, row in enumerate(inf1.iloc[:10].itertuples()): # E aqui preencho com os objetos que estão em cada coluna que ja preenchi
    table_row1 = table1.add_row().cells 
    for col in range(len(columns1)): 
        table_row1[col].text = str(row[col+1])
        
table1.alignment = WD_TABLE_ALIGNMENT.CENTER # Alinhamento das tabelas
table1.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


doc.save("Contrato pronto.docx")